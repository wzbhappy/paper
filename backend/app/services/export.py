"""稿件组装、引用一致性校验与多格式导出。

导出优先级（项目决策）：Markdown > LaTeX > Word。
Word 用 python-docx 生成，未安装时明确报错而非静默降级。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from loguru import logger

from app.services.review import CITATION_RE, extract_citations

AI_DISCLOSURE = (
    "本文部分内容在人工智能辅助下生成，作者已对全部内容进行核查并承担最终责任。"
)


@dataclass
class ManuscriptPart:
    """一个待导出的章节。"""

    title: str
    level: int
    content: str = ""
    paper_ids: list[str] = field(default_factory=list)
    """该节引用的文献，顺序对应节内 [n] 编号。"""
    ai_generated: bool = False


@dataclass
class Reference:
    paper_id: str
    title: str | None = None
    authors: str | None = None
    year: int | None = None
    venue: str | None = None
    doi: str | None = None

    def citation_key(self) -> str:
        surname = "unknown"
        if self.authors:
            first = self.authors.split(",")[0].strip()
            if first:
                surname = re.sub(r"\W+", "", first.split()[-1]).lower() or "unknown"
        word = "untitled"
        if self.title:
            tokens = re.findall(r"\w+", self.title)
            if tokens:
                word = tokens[0].lower()
        return f"{surname}{self.year or 'nd'}{word}"


@dataclass
class CitationIssue:
    section: str
    kind: str
    """out_of_range / unused_reference / empty_section"""
    detail: str


@dataclass
class Manuscript:
    title: str
    parts: list[ManuscriptPart] = field(default_factory=list)
    references: list[Reference] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return sum(len(p.content) for p in self.parts)

    @property
    def has_ai_content(self) -> bool:
        return any(p.ai_generated for p in self.parts)


def _global_index(manuscript: Manuscript) -> dict[str, int]:
    return {ref.paper_id: i for i, ref in enumerate(manuscript.references, start=1)}


def _renumber(content: str, local_ids: list[str], global_map: dict[str, int]) -> str:
    """把节内局部编号 [n] 换成全局编号。越界编号已在写作阶段剥离，这里再兜一次。"""

    def replace(match: re.Match[str]) -> str:
        numbers = [int(n.strip()) for n in match.group(1).split(",")]
        mapped: list[int] = []
        for n in numbers:
            if 1 <= n <= len(local_ids):
                target = global_map.get(local_ids[n - 1])
                if target and target not in mapped:
                    mapped.append(target)
        return f"[{','.join(str(n) for n in sorted(mapped))}]" if mapped else ""

    return CITATION_RE.sub(replace, content)


def check_citations(manuscript: Manuscript) -> list[CitationIssue]:
    """校验引用一致性：越界引用、未被引用的参考文献、空章节。"""
    issues: list[CitationIssue] = []
    cited_paper_ids: set[str] = set()

    for part in manuscript.parts:
        if part.content.strip():
            for number in extract_citations(part.content):
                if 1 <= number <= len(part.paper_ids):
                    cited_paper_ids.add(part.paper_ids[number - 1])
                else:
                    issues.append(
                        CitationIssue(
                            section=part.title,
                            kind="out_of_range",
                            detail=f"引用编号 [{number}] 超出本节可引用文献数（{len(part.paper_ids)}）",
                        )
                    )
        elif part.level <= 2:
            # 只对章级空节告警，子节留空常见于写作中途
            issues.append(
                CitationIssue(
                    section=part.title, kind="empty_section", detail="该章节尚无正文"
                )
            )

    for ref in manuscript.references:
        if ref.paper_id not in cited_paper_ids:
            issues.append(
                CitationIssue(
                    section="参考文献",
                    kind="unused_reference",
                    detail=f"《{ref.title or ref.paper_id}》列于参考文献但正文未引用",
                )
            )

    return issues


def to_markdown(manuscript: Manuscript, include_disclosure: bool = True) -> str:
    global_map = _global_index(manuscript)
    lines = [f"# {manuscript.title}", ""]

    for part in manuscript.parts:
        lines.append(f"{'#' * min(part.level + 1, 6)} {part.title}")
        lines.append("")
        if part.content.strip():
            lines.append(_renumber(part.content, part.paper_ids, global_map))
            lines.append("")

    if manuscript.references:
        lines.append("## 参考文献")
        lines.append("")
        for i, ref in enumerate(manuscript.references, start=1):
            parts = [f"[{i}]"]
            if ref.authors:
                parts.append(f"{ref.authors}.")
            parts.append(f"{ref.title or ref.paper_id}.")
            if ref.venue:
                parts.append(f"{ref.venue}.")
            if ref.year:
                parts.append(f"{ref.year}.")
            if ref.doi:
                parts.append(f"DOI: {ref.doi}.")
            lines.append(" ".join(parts))
        lines.append("")

    if include_disclosure and manuscript.has_ai_content:
        lines.append("---")
        lines.append("")
        lines.append(f"*{AI_DISCLOSURE}*")
        lines.append("")

    return "\n".join(lines)


_LATEX_ESCAPES = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}

# 单次扫描替换。顺序替换会出错：先把 \ 换成 \textbackslash{}，
# 随后的花括号规则又会转义其中的 {}，产生 \textbackslash\{\}。
_LATEX_PATTERN = re.compile("|".join(re.escape(c) for c in _LATEX_ESCAPES))

LATEX_SECTION_CMD = {1: "section", 2: "subsection", 3: "subsubsection"}


def escape_latex(text: str) -> str:
    """转义 LaTeX 特殊字符（单次扫描，避免二次转义）。"""
    return _LATEX_PATTERN.sub(lambda m: _LATEX_ESCAPES[m.group(0)], text)


def to_latex(manuscript: Manuscript, include_disclosure: bool = True) -> str:
    """生成可直接编译的 LaTeX。引用转成 \\cite{key}，参考文献用 thebibliography。"""
    global_map = _global_index(manuscript)
    key_by_index = {
        i: ref.citation_key() for i, ref in enumerate(manuscript.references, start=1)
    }

    def to_cite(content: str, local_ids: list[str]) -> str:
        def replace(match: re.Match[str]) -> str:
            numbers = [int(n.strip()) for n in match.group(1).split(",")]
            keys: list[str] = []
            for n in numbers:
                if 1 <= n <= len(local_ids):
                    index = global_map.get(local_ids[n - 1])
                    if index and key_by_index.get(index) not in keys:
                        keys.append(key_by_index[index])
            return f"\\cite{{{','.join(keys)}}}" if keys else ""

        return CITATION_RE.sub(replace, content)

    lines = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{ctex}",
        r"\usepackage{graphicx}",
        r"\usepackage{amsmath}",
        r"\usepackage{hyperref}",
        "",
        f"\\title{{{escape_latex(manuscript.title)}}}",
        r"\begin{document}",
        r"\maketitle",
        "",
    ]

    for part in manuscript.parts:
        command = LATEX_SECTION_CMD.get(part.level, "paragraph")
        lines.append(f"\\{command}{{{escape_latex(part.title)}}}")
        if part.content.strip():
            # 先转义再替换引用，避免 \cite 的花括号被转义
            body = escape_latex(part.content)
            # 转义后 [1] 形式不受影响，可直接匹配
            lines.append(to_cite(body, part.paper_ids))
        lines.append("")

    if manuscript.references:
        width = str(len(manuscript.references))
        lines.append(f"\\begin{{thebibliography}}{{{width}}}")
        for ref in manuscript.references:
            pieces = []
            if ref.authors:
                pieces.append(escape_latex(ref.authors))
            pieces.append(escape_latex(ref.title or ref.paper_id))
            if ref.venue:
                pieces.append(escape_latex(ref.venue))
            if ref.year:
                pieces.append(str(ref.year))
            lines.append(f"\\bibitem{{{ref.citation_key()}}} " + ". ".join(pieces) + ".")
        lines.append(r"\end{thebibliography}")
        lines.append("")

    if include_disclosure and manuscript.has_ai_content:
        lines.append(r"\vspace{1em}")
        lines.append(r"\noindent\textit{" + escape_latex(AI_DISCLOSURE) + "}")
        lines.append("")

    lines.append(r"\end{document}")
    return "\n".join(lines)


def to_bibtex(manuscript: Manuscript) -> str:
    entries: list[str] = []
    for ref in manuscript.references:
        fields = [f"  title = {{{ref.title or 'untitled'}}}"]
        if ref.authors:
            fields.append(f"  author = {{{ref.authors}}}")
        if ref.year:
            fields.append(f"  year = {{{ref.year}}}")
        if ref.venue:
            fields.append(f"  journal = {{{ref.venue}}}")
        if ref.doi:
            fields.append(f"  doi = {{{ref.doi}}}")
        entries.append(
            "@article{" + ref.citation_key() + ",\n" + ",\n".join(fields) + "\n}"
        )
    return "\n\n".join(entries)


def to_docx_bytes(manuscript: Manuscript, include_disclosure: bool = True) -> bytes:
    """生成 Word 文档字节流。需要 python-docx。"""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed; install it to enable Word export"
        ) from exc

    import io

    global_map = _global_index(manuscript)
    document = Document()
    document.add_heading(manuscript.title, level=0)

    for part in manuscript.parts:
        document.add_heading(part.title, level=min(part.level, 4))
        if part.content.strip():
            body = _renumber(part.content, part.paper_ids, global_map)
            for paragraph in body.split("\n\n"):
                text = paragraph.strip()
                if text:
                    document.add_paragraph(text)

    if manuscript.references:
        document.add_heading("参考文献", level=1)
        for i, ref in enumerate(manuscript.references, start=1):
            pieces = [f"[{i}]"]
            if ref.authors:
                pieces.append(f"{ref.authors}.")
            pieces.append(f"{ref.title or ref.paper_id}.")
            if ref.venue:
                pieces.append(f"{ref.venue}.")
            if ref.year:
                pieces.append(f"{ref.year}.")
            document.add_paragraph(" ".join(pieces))

    if include_disclosure and manuscript.has_ai_content:
        document.add_paragraph("")
        document.add_paragraph(AI_DISCLOSURE)

    buffer = io.BytesIO()
    document.save(buffer)
    logger.info("docx exported: {} parts", len(manuscript.parts))
    return buffer.getvalue()
